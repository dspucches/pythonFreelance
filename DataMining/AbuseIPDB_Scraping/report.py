# import necessary libraries
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import requests
from bs4 import BeautifulSoup

# Create a new document
document = Document()

# Add a title
title = document.add_heading('Web Scraping Results', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Get the data to be included in the document
url = 'https://www.example.com'
response = requests.get(url)
soup = BeautifulSoup(response.content, 'html.parser')

# Create a paragraph for each data point
for data_point in soup.find_all('div', {'class': 'data-point'}):
    paragraph = document.add_paragraph(data_point.text.strip())
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(8)
    paragraph_format.space_after = Pt(8)

# Save the document
document.save('web_scraping_results.docx')